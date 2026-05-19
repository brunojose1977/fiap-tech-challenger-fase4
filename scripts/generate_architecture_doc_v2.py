#!/usr/bin/env python3
"""Gera Documento de Arquitetura V2 (.docx e .pdf) em docs/."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

DOCS = Path(__file__).resolve().parents[1] / "docs"

SECTIONS = [
    (
        "1. Visão geral",
        "O projeto combina dois pipelines: (A) detecção de pose/violência com YOLOv8 em "
        "ECS Fargate e buckets S3 dedicados; (B) transcrição de conversas em vídeo/áudio com "
        "Amazon Transcribe, análise de risco com ChatGPT (OpenAI) e entrega de TXT + PDF no S3.",
    ),
    (
        "2. Pipeline aws-transcribe-audio-from-video-conversations",
        "Entrada: bucket transcribe-violence-input-fiap-posttech-iadevs-tcfase04. "
        "O serviço Amazon Transcribe extrai texto; arquivos transcribed-text-<nome>.txt são "
        "gravados no bucket transcribe-violence-output-fiap-posttech-iadevs-tcfase0. Em seguida "
        "o modelo gpt-5.4 classifica risco (segurança, integridade, ameaça, crime, mulher) e "
        "gera ChatGPT-5.4-avaliacao-conteudo-<nome>.pdf.",
    ),
    (
        "3. Componentes AWS",
        "S3 (entrada/saída Transcribe + buckets YOLO), Amazon Transcribe, ECS Fargate, ECR, "
        "IAM (task role com transcribe:* e S3), políticas de bucket para transcribe.amazonaws.com, "
        "CloudWatch Logs, GitHub OIDC.",
    ),
    (
        "4. Segurança",
        "Credenciais AWS via IAM roles (ECS/OIDC). OPENAI_API_KEY apenas em GitHub Secrets ou "
        "Secrets Manager. Buckets com SSE AES256, bloqueio de acesso público e versionamento.",
    ),
    (
        "5. CI/CD",
        "ci-cd.yml (qualidade + ECR), run-fargate.yml (YOLO), "
        "aws-transcribe-audio-from-video-conversations.yml (Transcribe + ChatGPT no runner Ubuntu).",
    ),
    (
        "6. Secrets GitHub",
        "AWS_ROLE_ARN, OPENAI_API_KEY; opcionalmente buckets se diferentes do padrão. "
        "Para YOLO: ECR_REPOSITORY_NAME, ECS_* conforme README.",
    ),
]


def build_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Documento de Arquitetura V2", 0)
    doc.add_paragraph(
        "Projeto: YOLOv8 Pose + Amazon Transcribe + ChatGPT — FIAP PostTech IA Devs."
    )
    for title, body in SECTIONS:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(body)
        p.style.font.size = Pt(11)
    doc.save(path)


def build_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    story = [
        Paragraph("<b>Documento de Arquitetura V2</b>", styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            "YOLOv8 Pose + Amazon Transcribe + ChatGPT — infraestrutura e fluxos.",
            styles["Normal"],
        ),
        Spacer(1, 18),
    ]
    for title, body in SECTIONS:
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(body, styles["BodyText"]))
        story.append(Spacer(1, 14))
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(story)


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    docx_path = DOCS / "Documento de Arquitetura V2.docx"
    pdf_path = DOCS / "Documento de Arquitetura V2.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(f"Gerado: {docx_path}")
    print(f"Gerado: {pdf_path}")


if __name__ == "__main__":
    main()
